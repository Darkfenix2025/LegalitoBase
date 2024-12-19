import os
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cargar variables de entorno
load_dotenv()

# Configurar Gemini API Key
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Crear la configuración del modelo
generation_config = {
    "temperature": 0.6,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

# Crear modelo de Gemini
model = genai.GenerativeModel(
    model_name="gemini-2.0-flash-exp",
    generation_config=generation_config,
)

# Función para interactuar con Gemini
def gemini_query(text):
    prompt = f"""
    Eres un abogado profesional especializado en el sistema legal argentino. Tu principal tarea es reformular, corregir y estructurar los hechos proporcionados para que puedan ser utilizados como base en un escrito de demanda judicial. 
    Tu enfoque debe ser identificar los puntos clave de los hechos, especialmente aquellos que tienen relevancia jurídica, y justificar las consecuencias jurídicas en base al derecho aplicable.
    Utiliza un lenguaje claro, preciso y formal, adaptado al contexto judicial.

    Hechos proporcionados: "{text}"
    Reformula los hechos en un estilo formal y estructurado para incluir en una demanda judicial. Si es necesario, menciona normativa aplicable y explica las consecuencias jurídicas.
    """
    try:
        response = model.generate_content([prompt])
        return response.text
    except Exception as e:
        return f"Error al generar la respuesta: {str(e)}"

# Función para generar un documento de Word
def generar_documento(respuesta):
    try:
        doc = Document()

        # Título del documento
        titulo = doc.add_heading('Hechos con Encuadre Legal', level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Agregar párrafo de respuesta
        parrafo = doc.add_paragraph(respuesta)
        parrafo_format = parrafo.paragraph_format
        parrafo_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Guardar documento
        output_file = "hechos_reformulados.docx"
        doc.save(output_file)
        print(f"Documento generado exitosamente: {output_file}")
    except Exception as e:
        print(f"Error al generar el documento: {e}")

# Solicitar texto al usuario
text = input("Ingrese el texto a procesar: ")

# Llamar a Gemini y obtener la respuesta
respuesta = gemini_query(text)

# Verificar y generar el documento
if "Error" in respuesta:
    print(respuesta)
else:
    generar_documento(respuesta)
